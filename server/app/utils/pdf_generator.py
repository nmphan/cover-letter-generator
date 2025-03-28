"""
Utility functions for generating sample PDF files for testing
"""

from fpdf import FPDF


class PDF(FPDF):
    """A custom PDF class for generating sample documents."""

    def header(self):
        """Add a header to each page."""
        self.set_font("Arial", "B", 12)
        self.cell(0, 10, "Sample Job Description", 0, 1, "C")

    def footer(self):
        """Add a footer with page number to each page."""
        self.set_y(-15)
        self.set_font("Arial", "I", 8)
        self.cell(0, 10, f"Page {self.page_no()}", 0, 0, "C")

    def chapter_title(self, title):
        """Add a chapter title."""
        self.set_font("Arial", "B", 12)
        self.cell(0, 10, title, 0, 1, "L")
        self.ln(10)

    def chapter_body(self, body):
        """Add the main content of a chapter."""
        self.set_font("Arial", "", 12)
        self.multi_cell(0, 10, body)
        self.ln()


def create_sample_pdf(file_path: str) -> None:
    """
    Create a sample PDF file with a job description.

    Args:
        file_path: Path where the PDF file should be saved
    """
    pdf = PDF()
    pdf.add_page()
    pdf.chapter_title("Job Description")
    pdf.chapter_body(
        """
    We are looking for a skilled software developer to join our team. The ideal
    candidate should have experience with Python, SQL, and web development
    frameworks such as Django or Flask.

    Key responsibilities include:
    - Developing and maintaining web applications
    - Writing clean and efficient code
    - Collaborating with cross-functional teams
    - Troubleshooting and debugging issues

    Qualifications:
    - Bachelor's degree in Computer Science or related field
    - Strong knowledge of Python and SQL
    - Experience with web development frameworks
    - Excellent problem-solving skills
    - Good communication and teamwork skills
    """
    )
    pdf.output(file_path)


if __name__ == "__main__":
    create_sample_pdf("../sampleResume/sample_job_description.pdf")

# def generate_cover_letter_pdf(cover_letter_data, output_path):
#     """
#     Generate a PDF file containing a cover letter.
    
#     Args:
#         cover_letter_data (dict): Dictionary containing the cover letter content
#         output_path (str): Path where the PDF file should be saved
    
#     Returns:
#         str: Path to the generated PDF file
#     """
#     pdf = PDF()
#     pdf.add_page()
    
#     # Add sender info
#     pdf.set_font("Arial", "", 10)
#     pdf.cell(0, 5, cover_letter_data.get("sender_name", ""), 0, 1, "L")
#     pdf.cell(0, 5, cover_letter_data.get("sender_email", ""), 0, 1, "L")
#     pdf.cell(0, 5, cover_letter_data.get("sender_phone", ""), 0, 1, "L")
#     pdf.ln(5)
    
#     # Add date
#     pdf.cell(0, 5, cover_letter_data.get("date", ""), 0, 1, "L")
#     pdf.ln(5)
    
#     # Add recipient info
#     pdf.cell(0, 5, cover_letter_data.get("recipient_name", ""), 0, 1, "L")
#     pdf.cell(0, 5, cover_letter_data.get("recipient_company", ""), 0, 1, "L")
#     pdf.ln(10)
    
#     # Add salutation
#     pdf.cell(0, 5, cover_letter_data.get("salutation", "Dear Hiring Manager,"), 0, 1, "L")
#     pdf.ln(5)
    
#     # Add body paragraphs
#     pdf.set_font("Arial", "", 12)
#     body_text = cover_letter_data.get("body", "")
#     pdf.multi_cell(0, 10, body_text)
#     pdf.ln(5)
    
#     # Add closing
#     pdf.set_font("Arial", "", 10)
#     pdf.cell(0, 5, cover_letter_data.get("closing", "Sincerely,"), 0, 1, "L")
#     pdf.ln(15)
#     pdf.cell(0, 5, cover_letter_data.get("sender_name", ""), 0, 1, "L")
    
#     # Save the PDF
#     pdf.output(output_path)
#     return output_path

def generate_cover_letter_pdf(cover_letter_data, output_path):
    """
    Generate a PDF file containing a cover letter.
    
    Args:
        cover_letter_data (dict): Dictionary containing the cover letter content
        output_path (str): Path where the PDF file should be saved
    
    Returns:
        str: Path to the generated PDF file
    """
    pdf = PDF()
    pdf.add_page()
    
    # Extract data from the nested structure
    sender_info = cover_letter_data.get("sender_info", {})
    resume_data = cover_letter_data.get("resume_data", {})
    job_data = cover_letter_data.get("job_data", {})
    
    # Add sender info
    pdf.set_font("Arial", "", 10)
    pdf.cell(0, 5, sender_info.get("name", ""), 0, 1, "L")
    pdf.cell(0, 5, sender_info.get("email", ""), 0, 1, "L")
    pdf.cell(0, 5, sender_info.get("phone", ""), 0, 1, "L")
    pdf.ln(5)
    
    # Add date
    pdf.cell(0, 5, sender_info.get("date", ""), 0, 1, "L")
    pdf.ln(5)
    
    # Add recipient info (you might want to add these to your frontend data)
    pdf.cell(0, 5, "Hiring Manager", 0, 1, "L")
    pdf.cell(0, 5, job_data.get("company", "Company Name"), 0, 1, "L")
    pdf.ln(10)
    
    # Add salutation
    pdf.cell(0, 5, "Dear Hiring Manager,", 0, 1, "L")
    pdf.ln(5)
    
    # Generate body content from resume and job data
    body_text = f"""
    I'm excited to apply for the position at {job_data.get("company", "your company")}. 
    With my skills in {', '.join(resume_data.get('skills', []))}, I believe I'd be a great fit.
    
    My experience includes:
    - {resume_data.get('experience', [{}])[0].get('description', 'Relevant experience')}
    
    I'm particularly interested in this role because {job_data.get('description', 'it aligns with my skills')}.
    """
    
    # Add body paragraphs
    pdf.set_font("Arial", "", 12)
    pdf.multi_cell(0, 10, body_text)
    pdf.ln(5)
    
    # Add closing
    pdf.set_font("Arial", "", 10)
    pdf.cell(0, 5, "Sincerely,", 0, 1, "L")
    pdf.ln(15)
    pdf.cell(0, 5, sender_info.get("name", ""), 0, 1, "L")
    
    # Save the PDF
    pdf.output(output_path)
    return output_path

def generate_cover_letter_docx(cover_letter_data, output_path):
    """
    Generate a DOCX file containing a cover letter.
    
    Args:
        cover_letter_data (dict): Dictionary containing the cover letter content
        output_path (str): Path where the DOCX file should be saved
    
    Returns:
        str: Path to the generated DOCX file
    """
    from docx import Document
    from docx.shared import Pt
    
    document = Document()
    
    # Add sender info
    document.add_paragraph(cover_letter_data.get("sender_name", ""))
    document.add_paragraph(cover_letter_data.get("sender_email", ""))
    document.add_paragraph(cover_letter_data.get("sender_phone", ""))
    document.add_paragraph()
    
    # Add date
    document.add_paragraph(cover_letter_data.get("date", ""))
    document.add_paragraph()
    
    # Add recipient info
    document.add_paragraph(cover_letter_data.get("recipient_name", ""))
    document.add_paragraph(cover_letter_data.get("recipient_company", ""))
    document.add_paragraph()
    
    # Add salutation
    document.add_paragraph(cover_letter_data.get("salutation", "Dear Hiring Manager,"))
    document.add_paragraph()
    
    # Add body paragraphs
    body_text = cover_letter_data.get("body", "")
    document.add_paragraph(body_text)
    document.add_paragraph()
    
    # Add closing
    document.add_paragraph(cover_letter_data.get("closing", "Sincerely,"))
    document.add_paragraph()
    document.add_paragraph(cover_letter_data.get("sender_name", ""))
    
    # Save the document
    document.save(output_path)
    return output_path